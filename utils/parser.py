"""
Resume Parser
Supports PDF, DOCX, and plain TXT formats.
"""
import io

def parse_resume(data: bytes, ext: str) -> str:
    """Parse raw file bytes into plain text."""
    ext = ext.lower()
    if ext == 'txt':
        return _parse_txt(data)
    elif ext == 'pdf':
        return _parse_pdf(data)
    elif ext == 'docx':
        return _parse_docx(data)
    return ""

def _parse_txt(data: bytes) -> str:
    for enc in ('utf-8', 'latin-1', 'ascii'):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')

def _parse_pdf(data: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        # Fallback: try pdfminer
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            out = io.StringIO()
            extract_text_to_fp(io.BytesIO(data), out, laparams=LAParams())
            return out.getvalue()
        except ImportError:
            return "[PDF parsing requires PyPDF2 or pdfminer.six. Install with: pip install PyPDF2]"

def _parse_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        return "[DOCX parsing requires python-docx. Install with: pip install python-docx]"
